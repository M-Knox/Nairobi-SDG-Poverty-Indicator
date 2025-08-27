import pandas as pd
from config import poverty_rates, sdg_progress_thresholds

def generate_poverty_stats(results_df):
    """
    Generate poverty statistics from results dataframe using your actual columns.
    
    Args:
        results_df: DataFrame containing poverty analysis results
    
    Returns:
        Dictionary with statistical summaries for each poverty indicator
    """
    poverty_stats = {}
    
    # Map your actual column names to display names
    poverty_indicators = {
        'Overall Poverty Rate': 'subcounty_overall_poverty_rate',
        'Food Poverty Rate': 'subcounty_food_poverty_rate', 
        'Hardcore Poverty Rate': 'subcounty_hardcore_poverty_rate'
    }
    
    # Generate stats for each indicator
    for display_name, column_name in poverty_indicators.items():
        if column_name in results_df.columns:
            poverty_stats[display_name] = {
                'Mean': results_df[column_name].mean(),
                'Median': results_df[column_name].median(),
                'Min': results_df[column_name].min(),
                'Max': results_df[column_name].max(),
                'Std Dev': results_df[column_name].std(),
                'Q1': results_df[column_name].quantile(0.25),
                'Q3': results_df[column_name].quantile(0.75)
            }
    
    # Add reference values from your config for comparison
    if poverty_stats:
        # Add county reference values as a separate row in tables
        reference_values = {
            'County Reference': poverty_rates['overall_poverty'],
            'Food Poverty Ref': poverty_rates['food_poverty'], 
            'Hardcore Poverty Ref': poverty_rates['hardcore_poverty']
        }
    
    return poverty_stats

def generate_progress_counts(results_df):
    """
    Generate SDG progress counts using your existing sdg1_progress_category column.
    
    Args:
        results_df: DataFrame containing poverty analysis results with sdg1_progress_category
    
    Returns:
        Dictionary with SDG progress level counts
    """
    if 'sdg1_progress_category' in results_df.columns:
        # Use your existing progress categories
        return results_df['sdg1_progress_category'].value_counts().to_dict()
    
    # Fallback: if progress category doesn't exist, create from overall poverty rates
    if 'subcounty_overall_poverty_rate' in results_df.columns:
        progress_counts = {
            'Excellent Progress': 0,
            'Good Progress': 0, 
            'Moderate Progress': 0,
            'Slow Progress': 0,
            'Significant Challenges': 0
        }
        
        # Categorize each area based on your SDG thresholds
        for poverty_rate in results_df['subcounty_overall_poverty_rate']:
            if poverty_rate <= sdg_progress_thresholds['excellent']:
                progress_counts['Excellent Progress'] += 1
            elif poverty_rate <= sdg_progress_thresholds['good']:
                progress_counts['Good Progress'] += 1
            elif poverty_rate <= sdg_progress_thresholds['moderate']:
                progress_counts['Moderate Progress'] += 1
            elif poverty_rate <= sdg_progress_thresholds['slow']:
                progress_counts['Slow Progress'] += 1
            else:
                progress_counts['Significant Challenges'] += 1
        
        # Remove categories with zero counts for cleaner display
        progress_counts = {k: v for k, v in progress_counts.items() if v > 0}
        
        return progress_counts
    
    return {}

def generate_enhanced_summary_report(results, poverty_stats, progress_counts):
    """
    Enhanced summary report using your actual column names.
    """
    # Use your actual column names
    poverty_col = 'subcounty_overall_poverty_rate'
    admin_col = 'Administrative_Unit'
    
    # Calculate key statistics
    highest_poverty_area = results.loc[results[poverty_col].idxmax(), admin_col]
    highest_poverty_rate = results[poverty_col].max()
    lowest_poverty_area = results.loc[results[poverty_col].idxmin(), admin_col]
    lowest_poverty_rate = results[poverty_col].min()
    average_poverty_rate = results[poverty_col].mean()
    county_reference = poverty_rates['overall_poverty']
    
    # Calculate areas above/below county average
    above_county_avg = len(results[results[poverty_col] > county_reference])
    total_areas = len(results)
    total_population = results['Total_Population'].sum()
    estimated_poor = results['estimated_poor_population'].sum() if 'estimated_poor_population' in results.columns else 0
    
    summary_parts = []
    
    # Overview with population context
    summary_parts.append(
        f"This report analyzes poverty indicators across {total_areas} {results['Administrative_level'].iloc[0].lower()}s "
        f"in Nairobi County, covering a total population of {total_population:,.0f} people. "
        f"The analysis reveals subcounty poverty rates ranging from {lowest_poverty_rate:.1f}% "
        f"to {highest_poverty_rate:.1f}%, with an average of {average_poverty_rate:.1f}%. "
        f"For context, the overall Nairobi County poverty rate is {county_reference}% according to "
        f"the Kenya Poverty Report 2019."
    )
    
    # Poverty impact analysis
    if estimated_poor > 0:
        summary_parts.append(
            f"The analysis estimates approximately {estimated_poor:,.0f} people "
            f"({(estimated_poor/total_population)*100:.1f}% of the total population) are living in poverty. "
            f"Comparative analysis shows that {above_county_avg} out of {total_areas} subcounties "
            f"({(above_county_avg/total_areas)*100:.1f}%) have poverty rates above the county average. "
            f"{highest_poverty_area} shows the highest poverty rate at {highest_poverty_rate:.1f}%, "
            f"while {lowest_poverty_area} demonstrates the lowest rate at {lowest_poverty_rate:.1f}%."
        )
    else:
        summary_parts.append(
            f"Comparative analysis shows that {above_county_avg} out of {total_areas} subcounties "
            f"({(above_county_avg/total_areas)*100:.1f}%) have poverty rates above the county average. "
            f"{highest_poverty_area} shows the highest poverty rate at {highest_poverty_rate:.1f}%, "
            f"while {lowest_poverty_area} demonstrates the lowest rate at {lowest_poverty_rate:.1f}%."
        )
    
    # SDG progress context
    if progress_counts:
        best_performing = min(progress_counts.items(), key=lambda x: list(progress_counts.keys()).index(x[0]))
        worst_performing = max(progress_counts.items(), key=lambda x: list(progress_counts.keys()).index(x[0]))
        
        summary_parts.append(
            f"Regarding Sustainable Development Goal 1 (No Poverty) progress, the distribution shows "
            f"{best_performing[1]} subcounties with '{best_performing[0]}' status and "
            f"{worst_performing[1]} subcounties facing '{worst_performing[0]}'. "
            f"This analysis provides crucial insights for targeted policy interventions and "
            f"resource allocation to achieve SDG poverty reduction targets across Nairobi's administrative units."
        )
    
    return "\n\n".join(summary_parts)

# Usage example with your actual data:
def create_complete_poverty_report(results, reports_dir):
    """
    Complete function that generates all required data and creates the report.
    
    Args:
        results: Your DataFrame with poverty analysis results
        reports_dir: Path to reports directory from your config
    """
    
    # Generate the required statistics using your config
    poverty_stats = generate_poverty_stats(results)
    progress_counts = generate_progress_counts(results)
    
    # Use the enhanced summary that incorporates your config values
    summary_report = generate_enhanced_summary_report(results, poverty_stats, progress_counts)
    
    # Create the Word document (using the function from the previous artifact)
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()

    # Add title
    title = doc.add_heading('Nairobi County Poverty Indicators Summary Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add main summary content
    doc.add_paragraph(summary_report)

    # Add reference section
    doc.add_heading('Reference Data', level=1)
    ref_para = doc.add_paragraph()
    ref_para.add_run("Nairobi County Baseline (Kenya Poverty Report 2019):\n").bold = True
    ref_para.add_run(f"• Overall Poverty: {poverty_rates['overall_poverty']}%\n")
    ref_para.add_run(f"• Food Poverty: {poverty_rates['food_poverty']}%\n")
    ref_para.add_run(f"• Hardcore Poverty: {poverty_rates['hardcore_poverty']}%")

    # Add statistical summary section
    if poverty_stats:
        doc.add_heading('Statistical Summary', level=1)
        stats_table = doc.add_table(rows=1, cols=len(list(poverty_stats.values())[0]) + 1)
        stats_table.style = 'Table Grid'

        # Add headers
        header_cells = stats_table.rows[0].cells
        header_cells[0].text = 'Indicator'
        for i, stat_name in enumerate(list(poverty_stats.values())[0].keys()):
            header_cells[i + 1].text = stat_name

        # Add data rows
        for indicator, stats in poverty_stats.items():
            row_cells = stats_table.add_row().cells
            row_cells[0].text = indicator
            for i, (stat_name, value) in enumerate(stats.items()):
                row_cells[i + 1].text = f"{value:.1f}%"

    # Add SDG Progress section
    if progress_counts:
        doc.add_heading('SDG Progress Distribution', level=1)
        progress_table = doc.add_table(rows=1, cols=3)
        progress_table.style = 'Table Grid'

        # Headers
        progress_header = progress_table.rows[0].cells
        progress_header[0].text = 'Progress Level'
        progress_header[1].text = 'Number of Areas'
        progress_header[2].text = 'Poverty Rate Threshold'

        # Data with thresholds
        threshold_ranges = {
            'Excellent Progress': f"≤{sdg_progress_thresholds['excellent']}%",
            'Good Progress': f"≤{sdg_progress_thresholds['good']}%",
            'Moderate Progress': f"≤{sdg_progress_thresholds['moderate']}%",
            'Slow Progress': f"≤{sdg_progress_thresholds['slow']}%",
            'Significant Challenges': f">{sdg_progress_thresholds['slow']}%"
        }

        for progress, count in progress_counts.items():
            row = progress_table.add_row().cells
            row[0].text = progress
            row[1].text = str(count)
            row[2].text = threshold_ranges.get(progress, "N/A")

    # Save document
    report_file = reports_dir / "nairobi_poverty_indicators_summary_report.docx"
    doc.save(report_file)

    # Print key findings using actual column names
    poverty_col = 'subcounty_overall_poverty_rate'
    admin_col = 'Administrative_Unit'
    
    print(f"Complete summary report exported to: {report_file}")
    print(f"Report includes {len(poverty_stats)} poverty indicators and {len(progress_counts)} progress categories")
    
    print("\nKey Findings:")
    print(f"  • Highest poverty: {results.loc[results[poverty_col].idxmax(), admin_col]} ({results[poverty_col].max():.1f}%)")
    print(f"  • Lowest poverty: {results.loc[results[poverty_col].idxmin(), admin_col]} ({results[poverty_col].min():.1f}%)")
    print(f"  • Average poverty: {results[poverty_col].mean():.1f}%")
    print(f"  • County reference: {poverty_rates['overall_poverty']}%")
    if 'estimated_poor_population' in results.columns:
        print(f"  • Total estimated poor population: {results['estimated_poor_population'].sum():,.0f}")
    
    return report_file

# Usage options:

# Option 1: If calling from another script/notebook, pass results as parameter
# report_file = create_complete_poverty_report(results, reports_dir)

# Option 2: Load results from a saved file
def create_report_from_file(results_file_path, reports_dir):
    """
    Load results from file and create report
    
    Args:
        results_file_path: Path to saved results file (CSV, Excel, etc.)
        reports_dir: Path to reports directory
    """
    import pandas as pd
    
    # Load results based on file extension
    if str(results_file_path).endswith('.csv'):
        results = pd.read_csv(results_file_path)
    elif str(results_file_path).endswith(('.xlsx', '.xls')):
        results = pd.read_excel(results_file_path)
    else:
        raise ValueError("Unsupported file format. Use CSV or Excel.")
    
    return create_complete_poverty_report(results, reports_dir)

# Option 3: If you want to run this as a standalone script
if __name__ == "__main__":
    # You'll need to specify where to load results from
    from config import processed_data_dir, reports_dir
    import pandas as pd
    
    # Example: Load from a saved results file
    # results_file = processed_data_dir / "poverty_analysis_results.csv"  # Update with your actual file
    # if results_file.exists():
    #     results = pd.read_csv(results_file)
    #     report_file = create_complete_poverty_report(results, reports_dir)
    # else:
    #     print(f"Results file not found: {results_file}")
    #     print("Please save your results DataFrame first or run the analysis notebook.")